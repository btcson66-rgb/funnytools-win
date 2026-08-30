from __future__ import annotations

import io
from collections import defaultdict

import pdfplumber
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt
from PIL import Image

from .common import MAX_OCR_PAGES, open_pdf_reader, run_tesseract_png, validate_ocr_lang
from .pdf_render import render_pdf_page_png


def _line_groups(words: list[dict], tolerance: float = 3.5) -> list[list[dict]]:
    rows: list[list[dict]] = []
    for word in sorted(words, key=lambda w: (float(w.get("top", 0)), float(w.get("x0", 0)))):
        top = float(word.get("top", 0))
        target = None
        for row in rows:
            row_top = sum(float(w.get("top", 0)) for w in row) / len(row)
            if abs(row_top - top) <= tolerance:
                target = row
                break
        if target is None:
            target = []
            rows.append(target)
        target.append(word)
    for row in rows:
        row.sort(key=lambda w: float(w.get("x0", 0)))
    return rows


def _add_editable_text(docx: Document, page) -> int:
    try:
        words = page.extract_words(
            use_text_flow=True,
            keep_blank_chars=False,
            extra_attrs=["fontname", "size"],
        )
    except Exception:
        words = page.extract_words(use_text_flow=True, keep_blank_chars=False)

    if not words:
        text = page.extract_text() or ""
        for paragraph in [x.strip() for x in text.split("\n\n") if x.strip()]:
            docx.add_paragraph(paragraph)
        return 1 if text.strip() else 0

    added = 0
    for row in _line_groups(words):
        p = docx.add_paragraph()
        previous_x1 = None
        for word in row:
            x0 = float(word.get("x0", 0))
            if previous_x1 is not None and x0 - previous_x1 > 18:
                p.add_run("\t")
            elif previous_x1 is not None:
                p.add_run(" ")
            run = p.add_run(str(word.get("text", "")))
            size = word.get("size")
            if size:
                try:
                    run.font.size = Pt(min(max(float(size), 7), 28))
                except Exception:
                    pass
            font = str(word.get("fontname", "")).lower()
            if "bold" in font:
                run.bold = True
            if "italic" in font or "oblique" in font:
                run.italic = True
            previous_x1 = float(word.get("x1", x0))
        added += 1
    return added


def _add_page_images(docx: Document, reader, page_index: int, max_images: int = 6) -> int:
    added = 0
    try:
        images = list(reader.pages[page_index].images)
    except Exception:
        return 0

    for image_file in images[:max_images]:
        try:
            pil = image_file.image
            if pil.width < 80 or pil.height < 80:
                continue
            if pil.mode in ("RGBA", "LA"):
                rgba = pil.convert("RGBA")
                bg = Image.new("RGB", rgba.size, "white")
                bg.paste(rgba, mask=rgba.getchannel("A"))
                pil = bg
            elif pil.mode not in ("RGB", "L"):
                pil = pil.convert("RGB")
            b = io.BytesIO()
            pil.save(b, "JPEG", quality=90, optimize=True)
            b.seek(0)
            docx.add_picture(b, width=Inches(5.8))
            added += 1
        except Exception:
            continue
    return added


def pdf_to_docx(
    pdf_bytes: bytes,
    filename: str = "document.pdf",
    ocr_mode: str = "auto",
    ocr_lang: str = "eng",
    include_images: bool = True,
) -> tuple[bytes, dict]:
    reader = open_pdf_reader(pdf_bytes, filename)
    if ocr_mode not in {"auto", "force", "off"}:
        raise ValueError("ocr_mode must be auto, force, or off")
    validate_ocr_lang(ocr_lang)

    docx = Document()
    section = docx.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    stats = {
        "pages": len(reader.pages),
        "ocr_pages": 0,
        "text_pages": 0,
        "images_added": 0,
    }

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            extracted = (page.extract_text() or "").strip()
            use_ocr = ocr_mode == "force" or (ocr_mode == "auto" and len(extracted) < 20)

            if use_ocr:
                if stats["ocr_pages"] >= MAX_OCR_PAGES:
                    raise ValueError(f"OCR is limited to {MAX_OCR_PAGES} pages per job")
                png = render_pdf_page_png(pdf_bytes, i, dpi=220)
                ocr_text = run_tesseract_png(png, ocr_lang, output_format="txt", psm=6).decode(
                    "utf-8", errors="replace"
                ).strip()
                for line in [x.strip() for x in ocr_text.splitlines() if x.strip()]:
                    docx.add_paragraph(line)
                stats["ocr_pages"] += 1
            else:
                _add_editable_text(docx, page)
                stats["text_pages"] += 1
                if include_images:
                    stats["images_added"] += _add_page_images(docx, reader, i)

            if i < len(pdf.pages) - 1:
                docx.add_page_break()

    out = io.BytesIO()
    docx.save(out)
    return out.getvalue(), stats
